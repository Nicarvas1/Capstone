"""
Script para crear la plantilla del Informe de Evaluacion Psicopedagogica 2026.
"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from lxml import etree

AZUL = RGBColor(0x50, 0x50, 0xCC)

# ─── Helper: agrega o reutiliza un parrafo en la celda ───────────────────────
def agregar_tag(celda, tag):
    """Limpia la celda y pone la etiqueta {{ tag }}."""
    # Si no hay parrafos, agregamos uno
    if len(celda.paragraphs) == 0:
        p = celda.add_paragraph()
    else:
        # Limpiar runs de todos los parrafos
        for p in celda.paragraphs:
            for run in p.runs:
                run.text = ""
        p = celda.paragraphs[0]
    run = p.add_run("{{ " + tag + " }}")
    run.font.color.rgb = AZUL
    return run

def agregar_tag_inline(celda, tag):
    """Agrega la etiqueta al FINAL del contenido existente (preserva el label)."""
    if len(celda.paragraphs) == 0:
        p = celda.add_paragraph()
    else:
        p = celda.paragraphs[-1]
    run = p.add_run(" {{ " + tag + " }}")
    run.font.color.rgb = AZUL

def limpiar_y_construir(celda, lineas):
    """Limpia la celda y escribe una lista de (texto, tag_o_None, bold)."""
    for p in celda.paragraphs:
        for run in p.runs:
            run.text = ""
    primer_p = celda.paragraphs[0] if celda.paragraphs else celda.add_paragraph()
    for run in primer_p.runs:
        run.text = ""
    primer_usado = False
    for (texto, tag, bold) in lineas:
        if not primer_usado:
            p = primer_p
            primer_usado = True
        else:
            p = celda.add_paragraph()
        if texto:
            r = p.add_run(texto)
            r.bold = bold
        if tag:
            r2 = p.add_run("{{ " + tag + " }}")
            r2.font.color.rgb = AZUL

# ─── Cargar documento original ───────────────────────────────────────────────
carpeta = "documentos colegios"
ruta_original = None
for a in os.listdir(carpeta):
    if "Plantilla" in a and "Psico" in a:
        ruta_original = os.path.join(carpeta, a)
        break

if not ruta_original:
    print("ERROR: No se encontro el documento original.")
    sys.exit(1)

print(f"Cargando: {ruta_original}")
doc = Document(ruta_original)

# ─── TABLA 0: Identificacion del estudiante ──────────────────────────────────
# F1 C2-4: Nombre de identidad
# F2 C2-4: Nombre social
# F3 C2-4: Fecha nacimiento
# F4 C2-4: Edad
# F5 C2-4: Establecimiento
# F6 C2-4: Curso (celdas 2,3,4 tienen 0 parrafos → agregar_tag los crea)
# F7 C2-4: Fecha evaluacion
# F8 C1-4: Diagnostico
# F9 C2-4: Fecha emision diagnostico
tabla0 = doc.tables[0]
agregar_tag(tabla0.rows[1].cells[2], "nombre_est_psico")
agregar_tag(tabla0.rows[2].cells[2], "nombre_social_est_psico")
agregar_tag(tabla0.rows[3].cells[2], "fecha_nac_psico")
agregar_tag(tabla0.rows[4].cells[2], "edad_psico")
agregar_tag(tabla0.rows[5].cells[2], "establecimiento_psico")
agregar_tag(tabla0.rows[6].cells[2], "curso_psico")
agregar_tag(tabla0.rows[7].cells[2], "fecha_evaluacion_psico")
agregar_tag(tabla0.rows[8].cells[1], "diagnostico_psico")
agregar_tag(tabla0.rows[9].cells[2], "fecha_emision_diag_psico")
print("OK Tabla 0 (Identificacion)")

# ─── TABLA 1: Motivo + Instrumentos ─────────────────────────────────────────
tabla1 = doc.tables[1]
agregar_tag_inline(tabla1.rows[1].cells[0], "check_ingreso_psico")
agregar_tag_inline(tabla1.rows[1].cells[1], "check_reevaluacion_psico")
agregar_tag_inline(tabla1.rows[1].cells[2], "check_otro_psico")
agregar_tag(tabla1.rows[2].cells[1], "instrumentos_psico")
print("OK Tabla 1 (Motivo + Instrumentos)")

# ─── TABLA 2: Historia escolar ───────────────────────────────────────────────
tabla2 = doc.tables[2]
agregar_tag(tabla2.rows[1].cells[0], "historia_escolar_psico")
print("OK Tabla 2 (Historia escolar)")

# ─── TABLA 3: Analisis cualitativo ──────────────────────────────────────────
tabla3 = doc.tables[3]

limpiar_y_construir(tabla3.rows[1].cells[0], [
    ("a)  Habilidades Cognitivas y Comunicativas\n\n", None, True),
    ("Respecto al desarrollo de las habilidades cognitivas, presenta un adecuado desarrollo en:\n", None, False),
    (None, "analisis_cog_fortalezas", False),
    ("\nPor otra parte, las habilidades cognitivas con menor desarrollo estan relacionadas con:\n", None, False),
    (None, "analisis_cog_dificultades", False),
    ("\nRespecto al desarrollo de las habilidades comunicativas, presenta un adecuado desarrollo en:\n", None, False),
    (None, "analisis_com_fortalezas", False),
    ("\nPor otra parte, las habilidades comunicativas con menor desarrollo estan relacionadas con:\n", None, False),
    (None, "analisis_com_dificultades", False),
])

limpiar_y_construir(tabla3.rows[2].cells[0], [
    ("b) Habilidades Personales, Socioemocionales y de Aproximacion al Aprendizaje\n\n", None, True),
    ("Respecto al desarrollo de las habilidades Personales y Socioemocionales, presenta un adecuado desarrollo en:\n", None, False),
    (None, "analisis_soc_fortalezas", False),
    ("\nPor otra parte, las habilidades sociales con menor desarrollo estan relacionadas con:\n", None, False),
    (None, "analisis_soc_dificultades", False),
    ("\nRespecto a las habilidades de Aproximacion al Aprendizaje, presenta un adecuado desarrollo en:\n", None, False),
    (None, "analisis_apr_fortalezas", False),
    ("\nPor otra parte, las habilidades de aproximacion al aprendizaje con menor desarrollo:\n", None, False),
    (None, "analisis_apr_dificultades", False),
])

limpiar_y_construir(tabla3.rows[3].cells[0], [
    ("c) Habilidades Motoras, de Autonomia y Sensoriales\n\n", None, True),
    ("Respecto al desarrollo de las habilidades Motoras, presenta un adecuado desarrollo en:\n", None, False),
    (None, "analisis_mot_fortalezas", False),
    ("\nPor otra parte, las habilidades motoras con menor desarrollo estan relacionadas con:\n", None, False),
    (None, "analisis_mot_dificultades", False),
    ("\nRespecto al desarrollo de Autonomia, presenta un adecuado desarrollo en:\n", None, False),
    (None, "analisis_aut_fortalezas", False),
    ("\nPor otra parte, la autonomia con menor desarrollo esta relacionada con:\n", None, False),
    (None, "analisis_aut_dificultades", False),
    ("\nRespecto al Procesamiento Sensorial, presenta un adecuado desarrollo en:\n", None, False),
    (None, "analisis_sen_fortalezas", False),
    ("\nPor otra parte, el procesamiento sensorial con menor desarrollo esta relacionado con:\n", None, False),
    (None, "analisis_sen_dificultades", False),
])
print("OK Tabla 3 (Analisis cualitativo)")

# ─── TABLA 4: Sintesis ───────────────────────────────────────────────────────
tabla4 = doc.tables[4]

limpiar_y_construir(tabla4.rows[1].cells[0], [
    ("a) Habilidades Cognitivas y Comunicativas\n\n", None, True),
    ("COGNITIVAS - Progresos / Necesidades a trabajar:\n", None, True),
    (None, "sint_cog_progresos", False),
    ("\nObstaculos / Desafios a superar:\n", None, False),
    (None, "sint_cog_obstaculos", False),
    ("\nFactores del entorno que influyen en el progreso:\n", None, False),
    (None, "sint_cog_factores", False),
    ("\nEstrategias que han funcionado:\n", None, False),
    (None, "sint_cog_estrategias", False),
    ("\nCOMUNICATIVAS - Necesidades a trabajar:\n", None, True),
    (None, "sint_com_necesidades", False),
    ("\nObstaculos / Desafios a superar:\n", None, False),
    (None, "sint_com_obstaculos", False),
])

limpiar_y_construir(tabla4.rows[2].cells[0], [
    ("b) Habilidades Personales, Socioemocionales y de Aproximacion al Aprendizaje\n\n", None, True),
    ("PERSONALES Y SOCIOEMOCIONALES - Aspectos positivos / Fortalezas:\n", None, True),
    (None, "sint_soc_fortalezas", False),
    ("\nNecesidades a trabajar:\n", None, False),
    (None, "sint_soc_necesidades", False),
    ("\nObstaculos / Desafios a superar:\n", None, False),
    (None, "sint_soc_obstaculos", False),
    ("\nEstrategias que han funcionado:\n", None, False),
    (None, "sint_soc_estrategias", False),
    ("\nAPROXIMACION AL APRENDIZAJE - Aspectos positivos / Fortalezas:\n", None, True),
    (None, "sint_apr_fortalezas", False),
    ("\nNecesidades a trabajar:\n", None, False),
    (None, "sint_apr_necesidades", False),
    ("\nObstaculos / Desafios a superar:\n", None, False),
    (None, "sint_apr_obstaculos", False),
    ("\nEstrategias que han funcionado:\n", None, False),
    (None, "sint_apr_estrategias", False),
])

limpiar_y_construir(tabla4.rows[3].cells[0], [
    ("c) Habilidades Motoras, de Autonomia y Sensoriales\n\n", None, True),
    ("MOTORAS - Necesidades a trabajar:\n", None, True),
    (None, "sint_mot_necesidades", False),
    ("\nObstaculos / Desafios a superar:\n", None, False),
    (None, "sint_mot_obstaculos", False),
    ("\nEstrategias que han funcionado:\n", None, False),
    (None, "sint_mot_estrategias", False),
    ("\nAUTONOMIA - Necesidades a trabajar:\n", None, True),
    (None, "sint_aut_necesidades", False),
    ("\nObstaculos / Desafios a superar:\n", None, False),
    (None, "sint_aut_obstaculos", False),
    ("\nEstrategias que han funcionado:\n", None, False),
    (None, "sint_aut_estrategias", False),
    ("\nPROCESAMIENTO SENSORIAL - Necesidades a trabajar:\n", None, True),
    (None, "sint_sen_necesidades", False),
    ("\nObstaculos / Desafios a superar:\n", None, False),
    (None, "sint_sen_obstaculos", False),
])
print("OK Tabla 4 (Sintesis)")

# ─── TABLA 5: Conclusion ────────────────────────────────────────────────────
tabla5 = doc.tables[5]
limpiar_y_construir(tabla5.rows[1].cells[0], [
    ("A partir de los resultados y pruebas aplicadas se desprenden las siguientes conclusiones:\n\n", None, False),
    ("Por area evaluada:\n", None, True),
    (None, "conclusion_areas", False),
    ("\nGenerales:\n", None, True),
    (None, "conclusion_general", False),
])
print("OK Tabla 5 (Conclusion)")

# ─── TABLA 6: Sugerencias ───────────────────────────────────────────────────
tabla6 = doc.tables[6]
for fila_idx, tag in [
    (1, "sug_establecimiento"),
    (2, "sug_equipo_aula"),
    (3, "sug_estudiante"),
    (4, "sug_familia"),
    (6, "sug_instituciones"),
]:
    celda = tabla6.rows[fila_idx].cells[0]
    if len(celda.paragraphs) == 0:
        p = celda.add_paragraph()
    else:
        p = celda.add_paragraph()
    r = p.add_run("{{ " + tag + " }}")
    r.font.color.rgb = AZUL
print("OK Tabla 6 (Sugerencias)")

# ─── TABLA 7: Profesional que emite el informe ──────────────────────────────
tabla7 = doc.tables[7]
for fila in tabla7.rows:
    if len(fila.cells) < 2:
        continue
    texto = fila.cells[0].text.strip().lower()
    tag = None
    if "nombre completo" in texto:
        tag = "nombre_prof_psico"
    elif texto == "rut":
        tag = "rut_prof_psico"
    elif "profesi" in texto:
        tag = "profesion_prof_psico"
    elif "registro" in texto:
        tag = "registro_prof_psico"
    if tag:
        agregar_tag(fila.cells[1], tag)
print("OK Tabla 7 (Profesional emisor)")

# ─── TABLA 10: Docente de aula ──────────────────────────────────────────────
tabla10 = doc.tables[10]
for fila in tabla10.rows:
    if len(fila.cells) < 2:
        continue
    texto = fila.cells[0].text.strip().lower()
    tag = None
    if "nombre completo" in texto:
        tag = "nombre_docente"
    elif texto == "rut":
        tag = "rut_docente"
    elif "profesi" in texto:
        tag = "profesion_docente"
    if tag:
        agregar_tag(fila.cells[1], tag)
print("OK Tabla 10 (Docente de aula)")

# ─── Guardar plantilla ──────────────────────────────────────────────────────
os.makedirs("planillas", exist_ok=True)
ruta_salida = "planillas/plantilla_psicopedagogica.docx"
doc.save(ruta_salida)
print(f"\nPlantilla guardada en: {ruta_salida}")
